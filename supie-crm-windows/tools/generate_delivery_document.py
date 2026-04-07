from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt
from playwright.sync_api import sync_playwright


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "20260403"
SCREEN_DIR = OUT_DIR / "delivery_assets"
DOCX_PATH = OUT_DIR / "20260403-软件系统交付文档编写方案.docx"
MD_PATH = OUT_DIR / "20260403-软件系统交付文档编写方案.md"
BASE_URL = "http://127.0.0.1:3000"

USERNAME = "admin"
PASSWORD = "admin123"
VIEWPORT = {"width": 1600, "height": 900}


@dataclass
class Shot:
    key: str
    title: str
    path: Path
    note: str


def ensure_dirs() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    SCREEN_DIR.mkdir(parents=True, exist_ok=True)


def login_and_capture() -> list[Shot]:
    shots: list[Shot] = []
    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True)
        context = browser.new_context(viewport=VIEWPORT, device_scale_factor=1)
        page = context.new_page()

        page.goto(f"{BASE_URL}/login", wait_until="networkidle")
        captcha = page.locator("span.login-captcha-box").inner_text().strip()
        page.fill('input[name="username"]', USERNAME)
        page.fill('input[name="password"]', PASSWORD)
        page.fill('input[name="captcha_input"]', captcha)
        page.get_by_role("button", name="登录").click()
        page.wait_for_url("**/workbench")
        page.wait_for_load_state("networkidle")

        page.goto(f"{BASE_URL}/projects", wait_until="networkidle")
        page.wait_for_timeout(800)
        first_project = page.locator('table tbody a[href^="/projects/"]').first
        project_href = first_project.get_attribute("href") or "/projects"

        capture_plan = [
            ("login", "登录页", "/login", "01-login.png", "用于快速入门与首次登录说明。"),
            ("workbench", "工作台", "/workbench", "02-workbench.png", "用于说明日常处理入口和待办节奏。"),
            ("customers", "客户管理", "/customers", "03-customers.png", "用于客户模块操作说明。"),
            ("project_detail", "项目详情", project_href, "04-project-detail.png", "用于项目交付、进展与结项说明。"),
            ("approvals", "审批中心", "/approvals", "05-approvals.png", "用于审批流程说明。"),
            ("users", "用户管理", "/users", "06-users.png", "用于管理员手册中的账号与角色配置说明。"),
            ("roles", "角色与权限", "/roles", "07-roles.png", "用于权限配置说明。"),
            ("invoices", "开票管理", "/invoices", "08-invoices.png", "用于财务使用说明。"),
        ]

        for key, title, route, filename, note in capture_plan:
            page.goto(f"{BASE_URL}{route}", wait_until="networkidle")
            page.wait_for_timeout(1000)
            out = SCREEN_DIR / filename
            page.screenshot(path=str(out), full_page=False)
            shots.append(Shot(key=key, title=title, path=out, note=note))

        browser.close()
    return shots


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_doc_style(doc: Document) -> None:
    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(10.5)
    for name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        if name in styles:
            styles[name].font.name = "Microsoft YaHei"
            styles[name]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def add_page_number_footer(section) -> None:
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(9)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "第 "
    r.append(t)
    fld.append(r)
    p._p.append(fld)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def add_table(doc: Document, headers: list[str], rows: list[list[str]], col_widths_cm: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    set_repeat_table_header(table.rows[0])
    for idx, header in enumerate(headers):
        hdr_cells[idx].text = header
        for p in hdr_cells[idx].paragraphs:
            for run in p.runs:
                run.font.bold = True
        set_cell_shading(hdr_cells[idx], "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    if col_widths_cm:
        for row in table.rows:
            for idx, width in enumerate(col_widths_cm):
                row.cells[idx].width = Cm(width)
    return table


def add_image(doc: Document, image_path: Path, caption: str, width_cm: float = 15.0) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = c.add_run(caption)
    r.italic = True
    r.font.size = Pt(9)


def build_markdown(shots: list[Shot]) -> str:
    lines: list[str] = []
    lines.append("# 软件系统交付文档编写方案")
    lines.append("")
    lines.append("## 一、文档目标")
    lines.append("本文件用于指导 `项目过程管理系统` 的交付文档编制工作，目的是把甲方验收、运维交接、用户使用、技术实现四类材料一次性说清楚，方便团队按岗位协作完成。")
    lines.append("")
    lines.append("## 二、建议分工")
    lines.append("| 类别 | 牵头角色 | 协作角色 | 主要职责 |")
    lines.append("| --- | --- | --- | --- |")
    lines.append("| 项目管理与验收类 | 项目经理 | 开发负责人、测试负责人、实施人员 | 统一交付口径、验收结论、变更版本、上线和回滚安排 |")
    lines.append("| 产品使用类 | 产品经理 / 测试负责人 | 开发负责人 | 编写用户手册、管理员手册、快速入门、FAQ，配截图说明 |")
    lines.append("| 技术部署与环境类 | 开发负责人 | 运维 / 实施人员 | 输出部署文档、安装步骤、环境配置、账号密钥清单 |")
    lines.append("| 开发与技术实现类 | 开发负责人 | 架构 / 后端 / 前端 | 输出架构、数据库、接口、编码规范说明 |")
    lines.append("")
    lines.append("## 三、交付文档清单")
    lines.append("| 文档名称 | 负责人 | 重点内容 | 是否建议截图 |")
    lines.append("| --- | --- | --- | --- |")
    docs = [
        ("项目交付说明书", "项目经理", "交付范围、系统版本、部署环境、交付清单总览", "建议"),
        ("项目验收报告", "项目经理", "功能、性能、安全验收结论，双方签字栏", "建议"),
        ("需求规格说明书", "产品经理", "最终确认版需求基线，和客户确认一致", "可选"),
        ("变更记录 / 版本说明", "项目经理", "需求变更、版本日志、更新内容、时间", "可选"),
        ("上线部署方案 + 回滚方案", "开发负责人", "上线步骤、责任人、时间节点、回滚步骤、数据保障", "不强制"),
        ("用户手册 / 操作手册", "产品经理 / 测试负责人", "功能步骤、界面元素、常见业务流程", "必须"),
        ("管理员手册", "测试负责人 / 开发负责人", "账号权限、后台配置、监控方法、操作说明", "建议"),
        ("快速入门指南", "产品经理", "最短上手路径，适合首次使用者", "必须"),
        ("FAQ / 常见问题排查", "测试负责人", "常见问题、排查方法、解决方案", "建议"),
        ("部署文档", "开发负责人", "服务器、中间件、数据库、端口、依赖清单", "不强制"),
        ("安装部署步骤", "开发负责人 / 实施人员", "自动化脚本 + 手动部署步骤", "不强制"),
        ("环境配置清单", "运维 / 实施人员", "生产、测试、备份环境参数", "不强制"),
        ("域名 / 证书 / 密钥 / 账号密码清单", "运维 / 实施人员", "安全交接信息、用途、有效期、责任人", "不强制"),
        ("系统架构设计文档", "开发负责人", "总体架构、模块划分、技术栈", "不强制"),
        ("数据库设计文档", "开发负责人", "ER 图、表结构、索引、SQL 脚本", "不强制"),
        ("接口文档", "开发负责人", "接口清单、入参出参、鉴权、调用示例", "不强制"),
        ("代码注释 / 规范说明", "开发负责人", "注释规范、编码规范、关键逻辑说明", "不强制"),
    ]
    for item in docs:
        lines.append(f"| {item[0]} | {item[1]} | {item[2]} | {item[3]} |")
    lines.append("")
    lines.append("## 四、各类文档怎么写")
    sections = {
        "项目管理与验收类": [
            "项目交付说明书：先把交付范围、系统版本、部署环境、交付清单列清楚，避免客户不知道“收到的是什么”。",
            "项目验收报告：验收结论要写明白，功能、性能、安全三个维度都要有结论，最后预留双方签字栏。",
            "需求规格说明书：只保留最终确认版，别把讨论中的版本混进去，重点是“客户确认过的需求是什么”。",
            "变更记录 / 版本说明：按时间线写，版本号、变更点、影响范围、确认人都要有。",
            "上线部署方案 + 回滚方案：上线怎么做、谁负责、什么时候做、出问题怎么退回，四件事都要写。",
        ],
        "产品使用类": [
            "用户手册 / 操作手册：要按普通人能看懂的方式写，最好按“先做什么、再做什么、看到什么结果”来组织。",
            "管理员手册：重点写账号、角色、系统配置、运行监控，方便交接给甲方管理员。",
            "快速入门指南：只保留最短的上手路径，适合第一次打开系统的人。",
            "FAQ / 常见问题排查：把最常见的问题写成“现象 + 原因 + 处理办法”的格式。",
        ],
        "技术部署与环境类": [
            "部署文档：把服务器配置、中间件版本、数据库信息、端口、依赖组件写完整，避免交接时口头补充。",
            "安装部署步骤：先给自动化脚本，再补手动步骤，步骤一定要按实际操作顺序写。",
            "环境配置清单：生产、测试、备份环境分别列开，不要混在一起。",
            "域名 / 证书 / 密钥 / 账号密码清单：属于高敏资料，建议加密交接并明确责任人。",
        ],
        "开发与技术实现类": [
            "系统架构设计文档：讲清楚系统怎么分层、模块怎么划分、技术栈怎么选。",
            "数据库设计文档：至少包括 ER 图、完整表结构、字段说明、索引、SQL 脚本。",
            "接口文档：接口清单、入参、出参、鉴权方式、示例都要有。",
            "代码注释 / 规范说明：如果交付源码，这份文档能减少后续维护成本。",
        ],
    }
    for title, items in sections.items():
        lines.append(f"### {title}")
        for item in items:
            lines.append(f"- {item}")
        lines.append("")
    lines.append("## 五、真实环境截图示例")
    lines.append("下面的截图来自真实浏览器环境，可直接放进用户手册、管理员手册和快速入门指南。")
    for shot in shots:
        lines.append(f"### {shot.title}")
        lines.append(f"![{shot.title}](delivery_assets/{shot.path.name})")
        lines.append("")
        lines.append(f"- 用途：{shot.note}")
        lines.append("")
    lines.append("## 六、建议交付顺序")
    lines.append("1. 先由项目经理把交付范围、验收口径、版本说明和部署安排定下来。")
    lines.append("2. 再由开发负责人补齐架构、数据库、接口和部署材料。")
    lines.append("3. 产品经理和测试人员一起写用户手册、快速入门、FAQ，并把截图放进去。")
    lines.append("4. 最后由项目经理统一校对命名、版本号、页序和交付清单。")
    lines.append("")
    lines.append("## 七、交付前自检")
    lines.append("- 文件名是否统一、是否能一眼看出用途。")
    lines.append("- 截图是否来自真实环境，页面标题是否清楚。")
    lines.append("- 每份文档是否都写了负责人、版本、日期和适用范围。")
    lines.append("- 用户手册和管理员手册是否足够像“给普通人看的说明书”。")
    lines.append("- 技术类文档是否把参数、路径、端口、账号等关键信息写全。")
    return "\n".join(lines)


def build_docx(shots: list[Shot]) -> None:
    doc = Document()
    set_doc_style(doc)
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    add_page_number_footer(section)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("软件系统交付文档编写方案")
    run.bold = True
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(20)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("适用于项目过程管理系统交付、运维、使用与技术说明")
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(10.5)

    doc.add_paragraph("")
    intro = doc.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = intro.add_run("说明：")
    run.bold = True
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run = intro.add_run("本文档按照岗位职责分工，把甲方交付时最需要的资料一次性拆清楚。写法尽量通俗，方便普通人阅读和后续交接。")
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    doc.add_heading("一、工作目标", level=1)
    doc.add_paragraph("本次交付文档的目标，不是“把文档写多”，而是把客户真正会用到、会验收、会接手、会排障的内容写完整。最终形成一套可以直接交付的材料清单，覆盖项目验收、用户使用、技术部署和技术实现四个层面。")

    doc.add_heading("二、建议分工", level=1)
    add_table(
        doc,
        ["类别", "牵头角色", "协作角色", "主要职责"],
        [
            ["项目管理与验收类", "项目经理", "开发负责人、测试负责人、实施人员", "统一交付口径、验收结论、变更版本、上线和回滚安排"],
            ["产品使用类", "产品经理 / 测试负责人", "开发负责人", "编写用户手册、管理员手册、快速入门、FAQ，配截图说明"],
            ["技术部署与环境类", "开发负责人", "运维 / 实施人员", "输出部署文档、安装步骤、环境配置、账号密钥清单"],
            ["开发与技术实现类", "开发负责人", "架构 / 后端 / 前端", "输出架构、数据库、接口、编码规范说明"],
        ],
        [3.5, 3.9, 4.1, 6.0],
    )
    doc.add_paragraph("如果团队人数不多，可以让一个人兼任多个角色，但文档责任一定要明确：谁写、谁审、谁最终交付，都要写清楚。")

    doc.add_heading("三、交付文档清单", level=1)
    add_table(
        doc,
        ["文档名称", "负责人", "重点内容", "是否建议截图"],
        [
            ["项目交付说明书", "项目经理", "交付范围、系统版本、部署环境、交付清单总览", "建议"],
            ["项目验收报告", "项目经理", "功能、性能、安全验收结论，双方签字栏", "建议"],
            ["需求规格说明书", "产品经理", "最终确认版需求基线，和客户确认一致", "可选"],
            ["变更记录 / 版本说明", "项目经理", "需求变更、版本日志、更新内容、时间", "可选"],
            ["上线部署方案 + 回滚方案", "开发负责人", "上线步骤、责任人、时间节点、回滚步骤、数据保障", "不强制"],
            ["用户手册 / 操作手册", "产品经理 / 测试负责人", "功能步骤、界面元素、常见业务流程", "必须"],
            ["管理员手册", "测试负责人 / 开发负责人", "账号权限、后台配置、监控方法、操作说明", "建议"],
            ["快速入门指南", "产品经理", "最短上手路径，适合首次使用者", "必须"],
            ["FAQ / 常见问题排查", "测试负责人", "常见问题、排查方法、解决方案", "建议"],
            ["部署文档", "开发负责人", "服务器、中间件、数据库、端口、依赖清单", "不强制"],
            ["安装部署步骤", "开发负责人 / 实施人员", "自动化脚本 + 手动部署步骤", "不强制"],
            ["环境配置清单", "运维 / 实施人员", "生产、测试、备份环境参数", "不强制"],
            ["域名 / 证书 / 密钥 / 账号密码清单", "运维 / 实施人员", "安全交接信息、用途、有效期、责任人", "不强制"],
            ["系统架构设计文档", "开发负责人", "总体架构、模块划分、技术栈", "不强制"],
            ["数据库设计文档", "开发负责人", "ER 图、表结构、索引、SQL 脚本", "不强制"],
            ["接口文档", "开发负责人", "接口清单、入参出参、鉴权、调用示例", "不强制"],
            ["代码注释 / 规范说明", "开发负责人", "注释规范、编码规范、关键逻辑说明", "不强制"],
        ],
        [5.0, 4.0, 6.0, 3.0],
    )

    doc.add_heading("四、各类文档怎么写", level=1)
    blocks = [
        ("项目管理与验收类", [
            "项目交付说明书：先把交付范围、系统版本、部署环境、交付清单列清楚，避免客户不知道“收到的是什么”。",
            "项目验收报告：验收结论要写明白，功能、性能、安全三个维度都要有结论，最后预留双方签字栏。",
            "需求规格说明书：只保留最终确认版，别把讨论中的版本混进去，重点是“客户确认过的需求是什么”。",
            "变更记录 / 版本说明：按时间线写，版本号、变更点、影响范围、确认人都要有。",
            "上线部署方案 + 回滚方案：上线怎么做、谁负责、什么时候做、出问题怎么退回，四件事都要写。",
        ]),
        ("产品使用类", [
            "用户手册 / 操作手册：要按普通人能看懂的方式写，最好按“先做什么、再做什么、看到什么结果”来组织。",
            "管理员手册：重点写账号、角色、系统配置、运行监控，方便交接给甲方管理员。",
            "快速入门指南：只保留最短的上手路径，适合第一次打开系统的人。",
            "FAQ / 常见问题排查：把最常见的问题写成“现象 + 原因 + 处理办法”的格式。",
        ]),
        ("技术部署与环境类", [
            "部署文档：把服务器配置、中间件版本、数据库信息、端口、依赖组件写完整，避免交接时口头补充。",
            "安装部署步骤：先给自动化脚本，再补手动步骤，步骤一定要按实际操作顺序写。",
            "环境配置清单：生产、测试、备份环境分别列开，不要混在一起。",
            "域名 / 证书 / 密钥 / 账号密码清单：属于高敏资料，建议加密交接并明确责任人。",
        ]),
        ("开发与技术实现类", [
            "系统架构设计文档：讲清楚系统怎么分层、模块怎么划分、技术栈怎么选。",
            "数据库设计文档：至少包括 ER 图、完整表结构、字段说明、索引、SQL 脚本。",
            "接口文档：接口清单、入参、出参、鉴权方式、示例都要有。",
            "代码注释 / 规范说明：如果交付源码，这份文档能减少后续维护成本。",
        ]),
    ]
    for title_text, bullet_items in blocks:
        doc.add_heading(title_text, level=2)
        add_bullets(doc, bullet_items)

    doc.add_heading("五、真实环境截图示例", level=1)
    doc.add_paragraph("下面这些截图都来自真实浏览器环境，可以直接放进用户手册、管理员手册和快速入门指南。")
    for shot in shots:
        doc.add_heading(shot.title, level=2)
        add_image(doc, shot.path, f"图：{shot.title}。{shot.note}")
        doc.add_paragraph(f"用途说明：{shot.note}")

    doc.add_heading("六、建议交付顺序", level=1)
    add_bullets(doc, [
        "先由项目经理把交付范围、验收口径、版本说明和部署安排定下来。",
        "再由开发负责人补齐架构、数据库、接口和部署材料。",
        "产品经理和测试人员一起写用户手册、快速入门、FAQ，并把截图放进去。",
        "最后由项目经理统一校对命名、版本号、页序和交付清单。",
    ])

    doc.add_heading("七、交付前自检", level=1)
    add_bullets(doc, [
        "文件名是否统一、是否能一眼看出用途。",
        "截图是否来自真实环境，页面标题是否清楚。",
        "每份文档是否都写了负责人、版本、日期和适用范围。",
        "用户手册和管理员手册是否足够像“给普通人看的说明书”。",
        "技术类文档是否把参数、路径、端口、账号等关键信息写全。",
    ])

    doc.save(DOCX_PATH)


def main() -> int:
    ensure_dirs()
    shots = login_and_capture()
    MD_PATH.write_text(build_markdown(shots), encoding="utf-8")
    build_docx(shots)
    print(MD_PATH)
    print(DOCX_PATH)
    print(SCREEN_DIR)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
