from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


def setup(doc: Document) -> None:
    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(10.5)
    for name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        if name in styles:
            styles[name].font.name = "Microsoft YaHei"
            styles[name]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def make_doc(title: str) -> Document:
    doc = Document()
    setup(doc)
    sec = doc.sections[0]
    sec.top_margin = Inches(0.7)
    sec.bottom_margin = Inches(0.65)
    sec.left_margin = Inches(0.8)
    sec.right_margin = Inches(0.8)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(18)
    return doc


def bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def numbered(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Number")


def add_code_block(doc: Document, lines: list[str]) -> None:
    for line in lines:
        p = doc.add_paragraph()
        p.style = doc.styles["Normal"]
        p.paragraph_format.left_indent = Inches(0.2)
        run = p.add_run(line)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(9.5)


def build() -> Document:
    doc = make_doc("安装部署步骤 - 详细版")
    doc.add_paragraph("本文档面向完全没有接触过本系统的人，目标是让你从零安装 Python、配置环境、修改数据库连接、启动系统并确认系统确实可用。")

    doc.add_heading("一、先说最重要的前提", level=1)
    bullet(doc, "推荐使用 Python 3.11.x。")
    bullet(doc, "最低可接受版本是 Python 3.10.x。")
    bullet(doc, "不要使用 Python 3.8 或更低版本。")
    bullet(doc, "数据库默认使用 PostgreSQL。")
    bullet(doc, "默认 Web 访问端口是 3000。")

    doc.add_heading("二、安装 Python 的完整步骤", level=1)
    numbered(doc, "打开浏览器，访问 Python 官方下载页。")
    numbered(doc, "下载 Windows 版 Python 3.11.x 安装包。")
    numbered(doc, "双击安装包，打开安装向导。")
    numbered(doc, "勾选 Add python.exe to PATH。")
    numbered(doc, "点击 Install Now。")
    numbered(doc, "等待安装完成，不要中途退出。")
    numbered(doc, "安装完成后关闭并重新打开 PowerShell。")
    numbered(doc, "输入 python --version。")
    numbered(doc, "确认输出为 Python 3.11.x。")
    numbered(doc, "如果版本不是 3.11.x，先调整系统 PATH，再继续。")

    doc.add_heading("三、创建虚拟环境", level=1)
    numbered(doc, "打开 PowerShell。")
    numbered(doc, "切换到项目目录：cd D:\\Supie\\crm。")
    numbered(doc, "执行 python -m venv .venv。")
    numbered(doc, "等待 .venv 目录生成。")
    numbered(doc, "执行 .\\.venv\\Scripts\\python --version，确认虚拟环境里的 Python 版本正确。")
    numbered(doc, "后续安装依赖和启动程序都在这个虚拟环境里执行。")

    doc.add_heading("四、Linux 系统下的安装与验证", level=1)
    doc.add_paragraph("如果你是在 Linux 服务器上部署，推荐按照下面顺序执行。这个流程和 Windows 的思路一样，区别只是命令不一样。")
    bullet(doc, "先确认系统里有 Python 3.11.x。")
    bullet(doc, "如果没有 Python 3.11.x，先安装系统依赖，再安装 Python。")
    bullet(doc, "建议使用系统自带的包管理器或者官方编译安装方式，不要随便拿来一个旧版本就直接用。")
    bullet(doc, "虚拟环境、依赖安装、数据库连接、启动验证，顺序不能乱。")

    doc.add_heading("Linux 安装 Python", level=2)
    numbered(doc, "执行 python3 --version，先看看系统里当前是什么版本。")
    numbered(doc, "如果版本低于 3.10，建议安装 Python 3.11。")
    numbered(doc, "在 Ubuntu / Debian 上，先安装常见编译依赖，例如 build-essential、libssl-dev、zlib1g-dev、libsqlite3-dev、libbz2-dev、libreadline-dev。")
    numbered(doc, "如果系统仓库里直接有 python3.11，可以优先使用系统包安装。")
    numbered(doc, "安装完成后再执行 python3.11 --version，确认安装成功。")
    numbered(doc, "如果 python3 指向的不是 3.11，可以在后续步骤里明确使用 python3.11。")

    doc.add_heading("Linux 创建虚拟环境", level=2)
    numbered(doc, "进入项目目录，例如 /workspace/crm 或你自己的实际项目路径。")
    numbered(doc, "执行 python3.11 -m venv .venv。")
    numbered(doc, "如果系统只支持 python3，可以先确认 python3 --version 是否已经是 3.11。")
    numbered(doc, "执行 .venv/bin/python --version，确认虚拟环境已经创建成功。")
    numbered(doc, "以后安装依赖和启动程序都建议使用这个虚拟环境里的 python。")

    doc.add_heading("Linux 安装依赖", level=2)
    numbered(doc, "在项目根目录执行 .venv/bin/pip install -r requirements.txt。")
    numbered(doc, "如果 pip 太旧，可以先执行 .venv/bin/python -m pip install --upgrade pip。")
    numbered(doc, "等待安装完成，不要中途退出。")
    numbered(doc, "如果有包安装失败，先记录完整报错，不要跳过。")
    numbered(doc, "如果服务器不能联网，需要先准备好离线 wheel 包，再进行安装。")

    doc.add_heading("Linux 修改数据库连接", level=2)
    doc.add_paragraph("Linux 下最稳妥的方式也是修改 .env 文件。做法和 Windows 一样，只是路径和编辑方式不同。")
    numbered(doc, "打开项目根目录的 .env 文件。")
    numbered(doc, "修改 PG_HOST、PG_PORT、PG_DATABASE、PG_USERNAME、PG_PASSWORD。")
    numbered(doc, "保存后重新启动应用。")
    numbered(doc, "如果你用的是 systemd 或其他服务方式，修改完后要重启服务，不只是重启终端。")
    doc.add_heading("Linux 启动系统", level=2)
    numbered(doc, "先确认虚拟环境和依赖都已经准备好。")
    numbered(doc, "在项目根目录执行 .venv/bin/python app.py。")
    numbered(doc, "如果项目提供了 service_runner 或 shell 启动脚本，也可以按项目规范执行。")
    numbered(doc, "确认控制台没有报数据库连接错误。")
    numbered(doc, "确认服务监听的端口和你配置的一致。")

    doc.add_heading("Linux 验证步骤", level=2)
    numbered(doc, "打开浏览器，访问 http://服务器IP:3000/login。")
    numbered(doc, "确认登录页能打开。")
    numbered(doc, "输入默认账号 admin 和默认密码 admin123。")
    numbered(doc, "确认可以登录并进入工作台。")
    numbered(doc, "检查客户、商机、合同、项目、审批、回款、开票页面是否都正常。")
    numbered(doc, "如果无法访问页面，先检查防火墙、端口开放情况和服务状态。")
    numbered(doc, "如果能打开页面但看不到数据，先检查数据库连接是否指向正确实例。")
    numbered(doc, "如果权限异常，检查管理员账号是否绑定了管理员角色。")

    doc.add_heading("Linux 常见问题", level=2)
    bullet(doc, "如果提示 python3.11 找不到，先确认系统里是否真的安装了 3.11。")
    bullet(doc, "如果提示 libssl、zlib、sqlite 之类缺失，说明系统编译依赖没有装全。")
    bullet(doc, "如果服务启动了但浏览器打不开，先检查防火墙和端口监听。")
    bullet(doc, "如果登录后页面空白，优先看日志和数据库连接。")
    bullet(doc, "如果你是用 SSH 远程登录服务器，不要把终端关掉，除非已经配置成服务。")

    doc.add_heading("五、安装项目依赖", level=1)
    numbered(doc, "保持当前目录仍然是项目根目录。")
    numbered(doc, "执行 .\\.venv\\Scripts\\pip install -r requirements.txt。")
    numbered(doc, "等待所有依赖安装完成。")
    numbered(doc, "如果提示 pip 过旧，先执行 .\\.venv\\Scripts\\python -m pip install --upgrade pip。")
    numbered(doc, "如果个别包安装失败，不要跳过，要把完整报错记下来。")
    numbered(doc, "如果安装时间较长，这是正常的。")

    doc.add_heading("六、如何验证 Python 和依赖", level=1)
    numbered(doc, "执行 python --version，确认是 3.11.x。")
    numbered(doc, "执行 .\\.venv\\Scripts\\python --version，确认虚拟环境生效。")
    numbered(doc, "执行 .\\.venv\\Scripts\\pip --version，确认 pip 可用。")
    numbered(doc, "执行一个简单导入检查，确认基础依赖能加载。")
    numbered(doc, "如果此处报错，不要急着启动系统，先把依赖问题解决。")

    doc.add_heading("七、如何修改数据库连接", level=1)
    doc.add_paragraph("系统优先读取项目根目录的 .env 文件。修改数据库连接，最推荐的方式是直接修改 .env。")
    numbered(doc, "打开项目根目录的 .env 文件。")
    numbered(doc, "找到 PG_HOST、PG_PORT、PG_DATABASE、PG_USERNAME、PG_PASSWORD。")
    numbered(doc, "把这些值改成你自己的数据库参数。")
    numbered(doc, "保存文件。")
    numbered(doc, "重启应用，新的数据库配置才会生效。")
    doc.add_paragraph("推荐的写法如下：")
    add_code_block(
        doc,
        [
            "PG_HOST=192.168.0.103",
            "PG_PORT=23083",
            "PG_DATABASE=zqq_test",
            "PG_USERNAME=zqq",
            "PG_PASSWORD=你的真实密码",
        ],
    )
    doc.add_paragraph("如果你只想临时测试，也可以在当前 PowerShell 窗口里设置环境变量：")
    add_code_block(
        doc,
        [
            '$env:PG_HOST=\"192.168.0.103\"',
            '$env:PG_PORT=\"23083\"',
            '$env:PG_DATABASE=\"zqq_test\"',
            '$env:PG_USERNAME=\"zqq\"',
            '$env:PG_PASSWORD=\"你的真实密码\"',
        ],
    )
    numbered(doc, "这种方式只对当前窗口有效，关闭窗口后就失效。")

    doc.add_heading("八、如何启动系统", level=1)
    numbered(doc, "先确认虚拟环境存在。")
    numbered(doc, "先确认依赖已经安装成功。")
    numbered(doc, "先确认数据库连接已经写对。")
    numbered(doc, "在 PowerShell 中进入项目目录。")
    numbered(doc, "执行 .\\.venv\\Scripts\\python app.py。")
    numbered(doc, "如果项目提供了 ops\\service_runner.py，也可以按项目规范使用它启动。")
    numbered(doc, "启动后，控制台应该能看到服务监听 3000 端口。")
    numbered(doc, "如果控制台报数据库错误，先回头检查 .env。")

    doc.add_heading("九、如何验证系统启动成功", level=1)
    numbered(doc, "打开浏览器，访问 http://127.0.0.1:3000/login。")
    numbered(doc, "确认登录页能打开。")
    numbered(doc, "输入默认账号 admin 和默认密码 admin123。")
    numbered(doc, "确认可以登录。")
    numbered(doc, "登录后先看工作台。")
    numbered(doc, "再继续看客户、商机、合同、项目、审批、回款、开票页面。")
    numbered(doc, "如果登录后空白或者报错，优先看控制台日志。")

    doc.add_heading("十、怎么判断数据库连接修改成功", level=1)
    numbered(doc, "修改完 .env 或环境变量后，重新启动程序。")
    numbered(doc, "登录页面能正常打开，说明 Web 服务没问题。")
    numbered(doc, "登录后页面有数据，说明数据库能连上。")
    numbered(doc, "如果页面空白但没有报错，检查是不是连接到了空库。")
    numbered(doc, "如果报连接失败，逐项核对主机、端口、库名、用户名、密码。")

    doc.add_heading("十一、常见错误与处理", level=1)
    bullet(doc, "如果提示 python 找不到，重新检查是否勾选了 Add python.exe to PATH。")
    bullet(doc, "如果提示依赖缺失，重新执行 pip install -r requirements.txt。")
    bullet(doc, "如果提示数据库连接失败，检查 PG_HOST、PG_PORT、PG_DATABASE、PG_USERNAME、PG_PASSWORD。")
    bullet(doc, "如果提示端口被占用，先关闭占用 3000 端口的程序。")
    bullet(doc, "如果你看到中文乱码，先确认你打开的是 DOCX，不要用记事本。")

    doc.add_heading("十二、建议的执行顺序", level=1)
    numbered(doc, "先检查 Python 版本。")
    numbered(doc, "再创建虚拟环境。")
    numbered(doc, "再安装依赖。")
    numbered(doc, "再修改数据库连接。")
    numbered(doc, "最后启动系统并验证登录。")

    return doc


def main() -> int:
    base = Path(__file__).resolve().parents[1]
    docs_root = base / "docs"
    accept_root = next(p for p in docs_root.iterdir() if p.is_dir() and any(c.is_dir() and c.name.startswith("03-") for c in p.iterdir()))
    tech_dir = next(c for c in accept_root.iterdir() if c.is_dir() and c.name.startswith("03-"))
    out = tech_dir / "install_deploy_clean.docx"
    build().save(out)
    print(out)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
