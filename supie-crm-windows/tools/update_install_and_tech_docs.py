from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]


def find_accept_root() -> Path:
    docs_root = ROOT / "docs"
    for child in docs_root.iterdir():
        if child.is_dir() and any(c.is_dir() and c.name.startswith("03-") for c in child.iterdir()):
            return child
    raise FileNotFoundError("找不到验收资料目录")


ROOT_DIR = find_accept_root()


def set_style(doc: Document) -> None:
    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(10.5)
    for name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        if name in styles:
            styles[name].font.name = "Microsoft YaHei"
            styles[name]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def set_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_footer(section) -> None:
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("验收资料").font.size = Pt(9)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        set_shading(table.rows[0].cells[i], "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = text


def bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def numbered(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Number")


def image(doc: Document, rel: str, width=Inches(5.8), caption: str | None = None) -> None:
    path = ROOT_DIR / "assets" / rel
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=width)
    if caption:
        c = doc.add_paragraph()
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.add_run(caption).italic = True


def save(doc: Document, path: Path) -> None:
    doc.save(path)


def make_doc(title: str) -> Document:
    doc = Document()
    set_style(doc)
    sec = doc.sections[0]
    sec.top_margin = Inches(0.7)
    sec.bottom_margin = Inches(0.65)
    sec.left_margin = Inches(0.8)
    sec.right_margin = Inches(0.8)
    add_footer(sec)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(18)
    return doc


def deployment_doc() -> Document:
    doc = make_doc("部署文档")
    doc.add_paragraph("本文件面向第一次接手部署的人，目标是让你按步骤做就能把系统跑起来。")
    doc.add_heading("一、部署前必须先确认的事情", level=1)
    bullet(doc, "你要知道本系统默认访问地址是 `http://127.0.0.1:3000`。")
    bullet(doc, "你要知道管理员默认账号是 `admin / admin123`，首次登录后建议立即修改。")
    bullet(doc, "你要知道系统现在依赖 Python 虚拟环境，不能直接裸跑。")
    bullet(doc, "你要知道数据库当前支持 PostgreSQL，环境变量没配对会直接连接失败。")
    bullet(doc, "你要知道部署前要先备份，尤其是数据库和上传文件。")

    doc.add_heading("二、环境准备清单", level=1)
    add_table(doc, ["项目", "最低要求", "说明"], [
        ["操作系统", "Windows 10 / Windows Server / Linux", "脚本是按这几种环境设计的"],
        ["Python", "3.10+", "虚拟环境建议单独创建"],
        ["数据库", "PostgreSQL", "默认环境按 PostgreSQL 连接"],
        ["端口", "3000", "如果被占用要改配置或先停其他服务"],
        ["浏览器", "Chrome / Edge", "建议用现代浏览器检查页面"],
    ])

    doc.add_heading("三、Windows 部署步骤", level=1)
    numbered(doc, "打开 PowerShell，并切到 `D:\\Supie\\crm`。")
    numbered(doc, "执行 `python -m venv .venv` 创建虚拟环境。")
    numbered(doc, r"执行 `.\.venv\Scripts\pip install -r requirements.txt` 安装依赖。")
    numbered(doc, r"执行 `.\.venv\Scripts\python app.py` 或 `ops\service_runner.py` 启动应用。")
    numbered(doc, "浏览器打开 `http://127.0.0.1:3000/login` 验证页面是否正常。")
    numbered(doc, "用默认账号登录，确认工作台、客户、项目和审批页面都能打开。")

    doc.add_heading("四、Linux 部署步骤", level=1)
    numbered(doc, "进入项目目录：`cd /workspace/crm`。")
    numbered(doc, "创建虚拟环境：`python3 -m venv .venv`。")
    numbered(doc, "安装依赖：`.venv/bin/pip install -r requirements.txt`。")
    numbered(doc, "启动应用：`.venv/bin/python app.py`。")
    numbered(doc, "访问登录页，确认服务已启动。")

    doc.add_heading("五、服务启动后应该检查什么", level=1)
    bullet(doc, "登录页是否能打开。")
    bullet(doc, "工作台是否有待办、关注事项和经营概览。")
    bullet(doc, "客户列表、项目列表、审批中心是否正常加载。")
    bullet(doc, "导出按钮是否可用。")
    bullet(doc, "用户、角色、权限页是否只对管理员开放。")

    doc.add_heading("六、常见部署失败原因", level=1)
    bullet(doc, "数据库账号密码写错。")
    bullet(doc, "端口 3000 被占用。")
    bullet(doc, "虚拟环境没创建好。")
    bullet(doc, "依赖没安装完成。")
    bullet(doc, "系统运行了，但浏览器打开的是错误地址。")

    doc.add_heading("七、如果部署失败，先按这个顺序排查", level=1)
    numbered(doc, "先确认服务进程是否真的在运行。")
    numbered(doc, "再确认浏览器访问的是不是 3000 端口。")
    numbered(doc, "然后看数据库连接配置。")
    numbered(doc, "最后看日志里有没有明显报错。")

    doc.add_heading("八、回滚与恢复", level=1)
    bullet(doc, "保留上一个可用版本代码。")
    bullet(doc, "保留数据库备份。")
    bullet(doc, "保留配置文件备份。")
    bullet(doc, "出问题时先回滚代码，再恢复数据库。")

    image(doc, "login.png", caption="登录页是部署检查的第一步。")
    image(doc, "workbench.png", caption="工作台是部署成功后第二个要检查的页面。")
    return doc


def install_steps_doc() -> Document:
    doc = make_doc("安装部署步骤")
    doc.add_paragraph("这份文档专门写给要亲手安装系统的人，按顺序做，别跳步。")
    doc.add_heading("一、安装前准备", level=1)
    bullet(doc, "准备一台能访问项目目录的机器。")
    bullet(doc, "准备 Python 环境。")
    bullet(doc, "确认数据库能连接。")
    bullet(doc, "准备好管理员账号、数据库账号和服务器权限。")

    doc.add_heading("二、Windows 手动安装", level=1)
    numbered(doc, "打开 PowerShell。")
    numbered(doc, "进入项目目录：`cd D:\\Supie\\crm`。")
    numbered(doc, "执行 `python -m venv .venv`。")
    numbered(doc, r"?? `\.venv\Scripts\pip install -r requirements.txt`?")
    numbered(doc, r"?? `\.venv\Scripts\python app.py` ???")
    numbered(doc, "浏览器访问 `http://127.0.0.1:3000/login`。")
    numbered(doc, "使用默认账号登录验证。")

    doc.add_heading("三、Linux 手动安装", level=1)
    numbered(doc, "进入项目目录。")
    numbered(doc, "创建虚拟环境。")
    numbered(doc, "安装依赖。")
    numbered(doc, "启动服务。")
    numbered(doc, "检查端口和页面。")

    doc.add_heading("四、安装后检查清单", level=1)
    bullet(doc, "登录页是否正常。")
    bullet(doc, "工作台是否正常。")
    bullet(doc, "客户列表是否正常。")
    bullet(doc, "项目列表和项目详情是否正常。")
    bullet(doc, "审批中心是否正常。")
    bullet(doc, "用户和角色管理是否管理员可见。")

    doc.add_heading("五、如果你安装不成功", level=1)
    bullet(doc, "先看命令有没有报错。")
    bullet(doc, "再看是不是虚拟环境没激活。")
    bullet(doc, "再看端口是不是被占用了。")
    bullet(doc, "再看数据库是不是没连上。")
    bullet(doc, "最后把报错信息截图给开发负责人。")
    return doc


def env_config_doc() -> Document:
    doc = make_doc("环境配置清单")
    doc.add_paragraph("本文件用于记录不同环境下的关键配置，避免上线、测试、备份时串环境。")
    doc.add_heading("一、环境划分", level=1)
    add_table(doc, ["环境", "用途", "注意事项"], [
        ["生产环境", "给实际用户使用", "配置最稳，变更最谨慎"],
        ["测试环境", "给测试和验收使用", "可提前验证版本"],
        ["备份环境", "给恢复和留档使用", "保存最关键的数据和配置"],
    ])

    doc.add_heading("二、关键配置项说明", level=1)
    bullet(doc, "`PG_HOST`：数据库主机地址。")
    bullet(doc, "`PG_PORT`：数据库端口。")
    bullet(doc, "`PG_DATABASE`：数据库名。")
    bullet(doc, "`PG_USERNAME`：数据库用户名。")
    bullet(doc, "`PG_PASSWORD`：数据库密码。")
    bullet(doc, "`APP_PORT`：Web 服务端口。")

    doc.add_heading("三、推荐的记录方式", level=1)
    add_table(doc, ["字段", "示例", "说明"], [
        ["环境名称", "生产", "写明到底是哪一套环境"],
        ["数据库地址", "192.168.0.103", "真实地址不要写错"],
        ["数据库端口", "23083", "与实际端口保持一致"],
        ["数据库名", "zqq_test", "每个环境要分开"],
        ["账号", "zqq", "和环境对应"],
        ["密码", "加密存储", "不要明文散落"],
    ])

    doc.add_heading("四、环境交接时的注意点", level=1)
    bullet(doc, "生产、测试、备份三个环境必须分开写。")
    bullet(doc, "账号密码不要混在普通使用文档里。")
    bullet(doc, "如果环境切换，先确认数据库和端口。")
    bullet(doc, "图片、附件、上传文件也要区分环境。")
    return doc


def architecture_doc() -> Document:
    doc = make_doc("系统架构设计文档")
    doc.add_paragraph("这份文档给技术人员看，重点是系统怎么分层、怎么流转、怎么维护。")
    doc.add_heading("一、系统整体架构", level=1)
    bullet(doc, "前端是 Jinja 模板页面，不是独立 SPA。")
    bullet(doc, "后端是 Flask 单体应用。")
    bullet(doc, "业务路由分在 `routes_system.py`、`routes_business.py`、`routes_projects.py`。")
    bullet(doc, "模板文件放在 `templates/`。")
    bullet(doc, "静态资源放在 `static/`。")
    bullet(doc, "数据默认走 PostgreSQL，部分环境也支持本地数据库。")

    doc.add_heading("二、主要模块", level=1)
    add_table(doc, ["模块", "说明", "典型页面"], [
        ["系统模块", "登录、工作台、用户、角色、权限", "登录页、工作台、用户管理、角色管理"],
        ["业务模块", "客户、商机、合同、回款、开票、审批", "列表页、详情页、审批中心"],
        ["项目模块", "项目、任务、里程碑、风险、进展", "项目列表、项目详情、子管理页"],
        ["支撑模块", "导出、附件、AI 辅助、日志", "导出接口、附件下载、AI 摘要"],
    ])

    doc.add_heading("三、核心数据流", level=1)
    bullet(doc, "用户先从登录入口进入系统。")
    bullet(doc, "工作台负责聚合待办和关注项。")
    bullet(doc, "客户、商机、合同之间形成销售链路。")
    bullet(doc, "合同之后进入项目交付链路。")
    bullet(doc, "项目交付又和回款、开票、审批连起来。")
    bullet(doc, "审批中心统一承接关键动作确认。")

    doc.add_heading("四、权限设计思路", level=1)
    numbered(doc, "先看菜单权限，决定看不看得到。")
    numbered(doc, "再看操作权限，决定能不能改、能不能批。")
    numbered(doc, "最后看数据权限，决定能看到多少范围。")

    doc.add_heading("五、页面组织方式", level=1)
    bullet(doc, "列表页负责筛选、导出和入口集中。")
    bullet(doc, "详情页负责展示一个对象的全部上下文。")
    bullet(doc, "工作台负责承接当天行动。")
    bullet(doc, "审批页负责处理流程确认。")

    doc.add_heading("六、后续扩展建议", level=1)
    bullet(doc, "如果以后要做多组织架构，可以继续往部门维度扩。")
    bullet(doc, "如果以后要做多项目组，可以在项目成员上继续扩。")
    bullet(doc, "如果以后要做更细的数据隔离，可以继续加数据权限规则。")

    return doc


def database_doc() -> Document:
    doc = make_doc("数据库设计文档")
    doc.add_paragraph("这份文档不是简单列表，而是告诉技术接手的人：每张表干什么、表之间什么关系、哪些字段最重要。")
    doc.add_heading("一、核心实体说明", level=1)
    add_table(doc, ["实体", "用途", "常见关系"], [
        ["users", "系统用户", "和 roles 通过 user_roles 关联"],
        ["roles", "角色定义", "被 users 和 permissions 使用"],
        ["customers", "客户台账", "关联 contacts、opportunities、contracts"],
        ["contacts", "联系人", "归属 customer"],
        ["opportunities", "商机", "关联 customer 和 contracts"],
        ["contracts", "合同", "关联 customer、receivables、invoices"],
        ["projects", "项目", "关联 customer、contracts、tasks 等"],
        ["tasks", "任务", "归属 project"],
        ["milestones", "里程碑", "归属 project"],
        ["risks", "风险", "归属 project"],
        ["receivables", "回款应收", "关联 contract"],
        ["invoices", "开票记录", "关联 contract 和 receivable"],
        ["approvals", "审批单", "关联商机、合同或项目"],
        ["attachments", "附件", "关联业务对象"],
        ["activity_logs", "操作日志", "记录系统动作"],
    ])

    doc.add_heading("二、每类表应该有什么字段", level=1)
    bullet(doc, "主键 `id`：每张表都应该有。")
    bullet(doc, "创建时间 `created_at`：记录对象何时产生。")
    bullet(doc, "更新时间 `updated_at`：记录对象最后一次修改。")
    bullet(doc, "状态字段 `status`：表示当前所处状态。")
    bullet(doc, "负责人字段 `owner` / `manager` / `applicant` / `approver`：记录归属人。")
    bullet(doc, "业务编号字段 `contract_no` / `invoice_code` 等：让人一眼识别。")

    doc.add_heading("三、关系理解方式", level=1)
    bullet(doc, "客户是前端经营的基础。")
    bullet(doc, "商机是客户推进后的阶段性成果。")
    bullet(doc, "合同是成交和履约的关键节点。")
    bullet(doc, "项目是交付执行的中心。")
    bullet(doc, "回款和开票是财务闭环。")
    bullet(doc, "审批是关键动作的确认机制。")

    doc.add_heading("四、数据库交付时要补什么", level=1)
    numbered(doc, "补真实 ER 图。")
    numbered(doc, "补完整字段说明。")
    numbered(doc, "补主键、外键和索引。")
    numbered(doc, "补初始化脚本和迁移脚本。")
    numbered(doc, "补典型查询示例。")

    doc.add_heading("五、排障时怎么用数据库文档", level=1)
    bullet(doc, "如果页面数据不对，先从表关系找。")
    bullet(doc, "如果某个列表查不到，先看筛选字段和状态字段。")
    bullet(doc, "如果权限异常，先看用户和角色表。")
    bullet(doc, "如果金额不对，先看合同、回款、开票之间的关联。")
    return doc


def api_doc() -> Document:
    doc = make_doc("接口文档")
    doc.add_paragraph("这份文档面向测试和对接人员，目标是让人一看就知道接口路径、用途、返回类型和调用时机。")
    doc.add_heading("一、接口分类", level=1)
    add_table(doc, ["模块", "方法", "路径", "说明"], [
        ["系统", "GET/POST", "/login", "登录"],
        ["系统", "GET", "/logout", "退出登录"],
        ["系统", "GET", "/workbench", "工作台"],
        ["系统", "GET", "/workbench/ai/priorities", "AI 优先建议"],
        ["用户", "GET/POST", "/users", "用户管理"],
        ["用户", "POST", "/users/<id>/toggle-active", "启停用户"],
        ["用户", "POST", "/users/<id>/reset-password", "重置密码"],
        ["用户", "POST", "/users/<id>/role", "配置角色"],
        ["角色", "GET/POST", "/roles", "角色列表与创建"],
        ["角色", "GET/POST", "/roles/<id>/edit", "编辑角色"],
        ["客户", "GET/POST", "/customers", "客户列表与新增"],
        ["客户", "GET", "/customers/<id>", "客户详情"],
        ["商机", "GET/POST", "/opportunities", "商机列表与新增"],
        ["商机", "GET", "/opportunities/<id>", "商机详情"],
        ["合同", "GET/POST", "/contracts", "合同列表与新增"],
        ["合同", "GET/POST", "/contracts/<id>", "合同详情与编辑"],
        ["项目", "GET/POST", "/projects", "项目列表与新增"],
        ["项目", "GET", "/projects/<id>", "项目详情"],
        ["项目", "POST", "/projects/<id>/progress", "新增进展"],
        ["项目", "POST", "/projects/<id>/submit-close-approval", "提交结项审批"],
        ["审批", "GET", "/approvals", "审批列表"],
        ["审批", "POST", "/approvals/<id>/approve", "同意审批"],
        ["审批", "POST", "/approvals/<id>/reject", "驳回审批"],
        ["审批", "POST", "/approvals/<id>/ai/summary", "审批 AI 摘要"],
        ["回款", "GET/POST", "/receivables", "回款应收"],
        ["开票", "GET/POST", "/invoices", "开票管理"],
    ])

    doc.add_heading("二、接口返回类型", level=1)
    bullet(doc, "页面接口通常返回 HTML。")
    bullet(doc, "AI 接口通常返回 JSON。")
    bullet(doc, "导出接口通常返回文件。")
    bullet(doc, "表单提交后大多会重定向回列表页或详情页。")

    doc.add_heading("三、测试时该怎么验接口", level=1)
    numbered(doc, "先确认登录状态。")
    numbered(doc, "再确认当前角色。")
    numbered(doc, "然后用正确参数调用。")
    numbered(doc, "最后检查返回值、页面状态和数据变化。")

    doc.add_heading("四、接口文档中还建议补什么", level=1)
    bullet(doc, "请求参数表。")
    bullet(doc, "返回值示例。")
    bullet(doc, "错误码说明。")
    bullet(doc, "鉴权说明。")
    bullet(doc, "调用示例。")

    doc.add_heading("五、对接时最容易出问题的地方", level=1)
    bullet(doc, "参数名写错。")
    bullet(doc, "路径写错。")
    bullet(doc, "忘了登录。")
    bullet(doc, "角色没权限。")
    bullet(doc, "返回类型和预期不一致。")
    return doc


def coding_doc() -> Document:
    doc = make_doc("代码注释 / 规范说明")
    doc.add_paragraph("这份文档是给接源码的人看的，目标是减少维护成本。")
    doc.add_heading("一、目录怎么分", level=1)
    bullet(doc, "系统路由放系统文件。")
    bullet(doc, "业务路由放业务文件。")
    bullet(doc, "项目路由放项目文件。")
    bullet(doc, "模板按页面类型放。")
    bullet(doc, "静态资源统一放到 `static/`。")

    doc.add_heading("二、命名怎么写", level=1)
    bullet(doc, "函数名要能看懂作用。")
    bullet(doc, "变量名不要太短。")
    bullet(doc, "文件名要和页面内容一致。")
    bullet(doc, "表单字段名要和后端参数尽量一致。")

    doc.add_heading("三、哪些地方必须写注释", level=1)
    numbered(doc, "权限判断。")
    numbered(doc, "审批逻辑。")
    numbered(doc, "复杂查询。")
    numbered(doc, "数据转换。")
    numbered(doc, "异常处理。")

    doc.add_heading("四、交付源码时建议附带什么", level=1)
    bullet(doc, "项目结构说明。")
    bullet(doc, "启动命令。")
    bullet(doc, "配置说明。")
    bullet(doc, "关键逻辑说明。")
    bullet(doc, "测试和回归命令。")

    doc.add_heading("五、接手代码的人最需要知道什么", level=1)
    bullet(doc, "先看页面入口。")
    bullet(doc, "再看业务路由。")
    bullet(doc, "再看模板。")
    bullet(doc, "最后看数据库和权限。")
    return doc


DOC_BUILDERS = {
    "03-技术部署与环境类/部署文档.docx": deployment_doc,
    "03-技术部署与环境类/安装部署步骤.docx": install_steps_doc,
    "03-技术部署与环境类/环境配置清单.docx": env_config_doc,
    "04-开发与技术实现类/系统架构设计文档.docx": architecture_doc,
    "04-开发与技术实现类/数据库设计文档.docx": database_doc,
    "04-开发与技术实现类/接口文档.docx": api_doc,
    "04-开发与技术实现类/代码注释与规范说明.docx": coding_doc,
}


def main() -> int:
    for rel, builder in DOC_BUILDERS.items():
        out = ROOT_DIR / rel
        out.parent.mkdir(parents=True, exist_ok=True)
        builder().save(out)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
